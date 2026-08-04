#!/usr/bin/env python3
"""Build the peer-blind English and Korean two-layer manuscripts as DOCX.

The builder uses the canonical SFTF Word template converted to DOCX as its
style/theme base.  It intentionally supports only the LaTeX constructs used by
draft/paper_en.tex and draft/paper_kr.tex so that the Word artifacts remain
deterministic and editable (real headings, equations, figures, and a real
table), rather than being a PDF-to-Word conversion.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TABLE_HELPER_DIR = Path(
    r"C:\Users\USER\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.802.11031\skills\documents\scripts"
)
sys.path.insert(0, str(TABLE_HELPER_DIR))
from table_geometry import (  # noqa: E402
    apply_table_geometry,
    column_widths_from_weights,
    section_content_width_dxa,
)

FIGURE_CAPTIONS = {
    "en": {
        "two_layer_workflow.png": (
            "Layer-first reconstruction. A single three-dimensional complex can "
            "contain cross-layer cells (A). The proposed route estimates separation "
            "and sampling adequacy from observed coordinates only (B). Accepted "
            "points are triangulated independently within each inferred layer, so a "
            "cross-layer face is not admissible (C)."
        ),
        "two_layer_results.png": (
            "Confirmatory outcomes generated directly from the frozen result "
            "artifacts. The method has the highest mean F-score and zero topology "
            "error in both panels. The lower row uses log10(error + 1) only for "
            "display; Table 1 reports the raw totals."
        ),
    },
    "kr": {
        "two_layer_workflow.png": (
            "Layer-first 재구성의 핵심. 하나의 3차원 complex는 두 층 사이의 "
            "잘못된 cell을 만들 수 있다(A). 관측 좌표만으로 층 분리와 sampling "
            "sufficiency를 검사한다(B). 통과한 경우 각 층을 따로 삼각분할하므로 "
            "층을 가로지르는 face는 허용되지 않는다(C)."
        ),
        "two_layer_results.png": (
            "동결 결과에서 직접 생성한 confirmatory 결과. 아래 topology panel은 "
            "표시를 위해 log10(error + 1)을 사용하며, 표 1에는 raw total을 제시했다."
        ),
    },
}


EQUATIONS = {
    "normal": "Q = (1/n) Σᵢ₌₁ⁿ nᵢnᵢᵀ,    n̂ = arg max‖v‖₌₁ vᵀQv.    (1)",
    "trend": "q(u,v) = β₀ + β₁u + β₂v + β₃u² + β₄uv + β₅v².    (2)",
    "snr": "SNR = |μ₁ − μ₀| / √[n⁻¹ Σᵢ(rᵢ − μ_zᵢ)²].    (3)",
    "construction": (
        "T_layer = D₂(Π₀X₀) ∪ D₂(Π₁X₁),    X₀ ∩ X₁ = ∅.    (4)"
    ),
}


RESULT_ROWS = [
    ["Synthetic (144)", "Layer-routed", "0.898536", "0.126043", "0"],
    ["", "B5 PCA-anisotropic alpha", "0.515603", "0.140552", "45,606"],
    ["", "M1 weighted power-alpha", "0.620140", "0.137152", "11,925"],
    ["S3DIS Area 5 (63)", "Layer-routed", "0.805611", "0.116385", "0"],
    ["", "B5 PCA-anisotropic alpha", "0.420983", "0.243356", "5,214"],
    ["", "M1 weighted power-alpha", "0.323764", "0.228920", "13,375"],
]


REF_MAP_EN = {
    "armeni2016": "Armeni et al. (2016)",
    "edelsbrunner1992": "Edelsbrunner (1992)",
    "edelsbrunner1994": "Edelsbrunner and Mücke (1994)",
    "teichmann1998": "Teichmann and Capps (1998)",
}
REF_MAP_KR = {
    "armeni2016": "[1]",
    "edelsbrunner1992": "[2]",
    "edelsbrunner1994": "[3]",
    "teichmann1998": "[4]",
}


def extract_braced(text: str, start: int) -> tuple[str, int]:
    """Extract balanced content when text[start] is an opening brace."""
    if start >= len(text) or text[start] != "{":
        raise ValueError("extract_braced must start on an opening brace")
    depth = 0
    for pos in range(start, len(text)):
        char = text[pos]
        if char == "{" and (pos == 0 or text[pos - 1] != "\\"):
            depth += 1
        elif char == "}" and (pos == 0 or text[pos - 1] != "\\"):
            depth -= 1
            if depth == 0:
                return text[start + 1 : pos], pos + 1
    raise ValueError("unbalanced braces")


def extract_command(text: str, command: str) -> str:
    match = re.search(rf"\\{re.escape(command)}\s*\{{", text)
    if not match:
        return ""
    return extract_braced(text, match.end() - 1)[0]


def strip_wrapping_commands(text: str) -> str:
    commands = (
        "textbf",
        "textit",
        "emph",
        "texttt",
        "mathbf",
        "mathsf",
        "mathbb",
        "mathcal",
        "mathrm",
        "operatorname",
        "widehat",
        "bar",
    )
    changed = True
    while changed:
        changed = False
        for command in commands:
            pattern = re.compile(rf"\\{command}\*?\s*\{{([^{{}}]*)\}}")
            text, count = pattern.subn(r"\1", text)
            changed = changed or count > 0
    return text


def clean_math(text: str) -> str:
    replacements = {
        r"\alpha": "α",
        r"\beta": "β",
        r"\mu": "μ",
        r"\ell": "ℓ",
        r"\Pi": "Π",
        r"\mathcal{D}": "D",
        r"\mathcal{T}": "T",
        r"\mathbb{R}": "ℝ",
        r"\varnothing": "∅",
        r"\emptyset": "∅",
        r"\subset": "⊂",
        r"\cup": "∪",
        r"\cap": "∩",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\in": "∈",
        r"\sum": "Σ",
        r"\frac": "",
        r"\sqrt": "√",
        r"\qquad": "    ",
        r"\quad": "  ",
        r"\,": " ",
        r"\!": "",
        r"\|": "‖",
        r"\circ": "°",
        r"\mathsf T": "T",
    }
    text = strip_wrapping_commands(text)
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace(r"\{", "{").replace(r"\}", "}")
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    text = re.sub(r"\\[A-Za-z]+\*?", "", text)
    return text.replace("{", "").replace("}", "")


def clean_latex(text: str, lang: str) -> str:
    refs = REF_MAP_EN if lang == "en" else REF_MAP_KR
    equation_numbers = {
        "eq:normal": "(1)",
        "eq:normal-kr": "(1)",
        "eq:trend": "(2)",
        "eq:trend-kr": "(2)",
        "eq:snr": "(3)",
        "eq:snr-kr": "(3)",
        "eq:construction": "(4)",
        "eq:construction-kr": "(4)",
    }
    figure_numbers = {
        "fig:workflow": "1",
        "fig:workflow-kr": "1",
        "fig:results": "2",
        "fig:results-kr": "2",
    }
    table_numbers = {"tab:results": "1", "tab:results-kr": "1"}

    def cite_repl(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        labels = [refs.get(key, key) for key in keys]
        if lang == "en" and match.group(0).startswith(r"\citep"):
            return "(" + "; ".join(labels) + ")"
        return "; ".join(labels)

    text = re.sub(r"\\citep?\{([^{}]+)\}", cite_repl, text)
    text = re.sub(
        r"\\eqref\{([^{}]+)\}",
        lambda m: equation_numbers.get(m.group(1), "(?)"),
        text,
    )
    text = re.sub(
        r"\\ref\{([^{}]+)\}",
        lambda m: figure_numbers.get(m.group(1), table_numbers.get(m.group(1), "?")),
        text,
    )
    text = re.sub(r"\$([^$]*)\$", lambda m: clean_math(m.group(1)), text)
    text = strip_wrapping_commands(text)
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = text.replace(r"\%", "%").replace(r"\_", "_")
    text = text.replace(r"\&", "&").replace(r"\#", "#")
    text = text.replace(r"M\"ucke", "Mücke")
    text = text.replace("--", "–").replace("~", " ")
    text = text.replace("``", "“").replace("''", "”")
    text = text.replace(r"\\", " ")
    text = re.sub(r"\\[A-Za-z]+\*?", "", text)
    text = text.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", text).strip()


def set_run_font(run, lang: str, size: float | None = None, bold: bool | None = None):
    font_name = "Times New Roman" if lang == "en" else "Malgun Gothic"
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), "Malgun Gothic" if lang == "kr" else font_name)


def set_style_font(style, lang: str, size: float | None = None):
    font_name = "Times New Roman" if lang == "en" else "Malgun Gothic"
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), "Malgun Gothic" if lang == "kr" else font_name)


def ensure_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_styles(doc: Document, lang: str) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, lang, 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.25)

    title = doc.styles["Title"]
    set_style_font(title, lang, 16)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, lang, 14)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.space_before = Pt(14)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.keep_with_next = True

    heading2 = doc.styles["Heading 2"]
    set_style_font(heading2, lang, 12)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(4)
    heading2.paragraph_format.keep_with_next = True

    abstract_title = ensure_style(doc, "Abstract Title")
    set_style_font(abstract_title, lang, 11)
    abstract_title.font.bold = True
    abstract_title.font.color.rgb = RGBColor(52, 90, 138)
    abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_title.paragraph_format.space_before = Pt(6)
    abstract_title.paragraph_format.space_after = Pt(4)

    abstract_body = ensure_style(doc, "Abstract Body")
    set_style_font(abstract_body, lang, 10)
    abstract_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_body.paragraph_format.line_spacing = 1.15
    abstract_body.paragraph_format.first_line_indent = Inches(0.25)
    abstract_body.paragraph_format.space_after = Pt(6)

    caption = doc.styles["Caption"]
    set_style_font(caption, lang, 10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = False

    equation = ensure_style(doc, "Equation")
    set_style_font(equation, lang, 11)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.line_spacing = 1.0
    equation.paragraph_format.space_before = Pt(6)
    equation.paragraph_format.space_after = Pt(6)
    equation.paragraph_format.keep_together = True

    bibliography = ensure_style(doc, "Bibliography")
    set_style_font(bibliography, lang, 10)
    bibliography.paragraph_format.left_indent = Inches(0.25)
    bibliography.paragraph_format.first_line_indent = Inches(-0.25)
    bibliography.paragraph_format.line_spacing = 1.15
    bibliography.paragraph_format.space_after = Pt(5)


def clear_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_header_footer(doc: Document, lang: str) -> None:
    short_title = (
        "Layer-First Two-Surface Reconstruction"
        if lang == "en"
        else "두 표면 점군의 layer-first 재구성"
    )
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(short_title)
        set_run_font(run, lang, 9)
        add_bottom_border(paragraph)

        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()
        add_page_number(footer_paragraph)


def add_text_paragraph(doc: Document, text: str, lang: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, lang)
    return paragraph


def add_caption(doc: Document, number: int, caption: str, lang: str) -> None:
    prefix = f"Figure {number}. " if lang == "en" else f"그림 {number}. "
    paragraph = doc.add_paragraph(style="Caption")
    run = paragraph.add_run(prefix)
    set_run_font(run, lang, 10, True)
    run = paragraph.add_run(caption)
    set_run_font(run, lang, 10)


def add_figure(doc: Document, image_path: Path, caption: str, number: int, lang: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(6.85))
    inline_shape._inline.docPr.set("title", f"Figure {number}")
    inline_shape._inline.docPr.set("descr", caption)
    add_caption(doc, number, caption, lang)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_results_table(doc: Document, lang: str) -> None:
    caption = (
        "Frozen confirmatory reconstruction results. Geometry loss and topology "
        "error are lower-is-better. All values are prespecified panel aggregates."
        if lang == "en"
        else "동결 confirmatory 결과. Geometry loss와 topology error는 낮을수록 좋다."
    )
    paragraph = doc.add_paragraph(style="Caption")
    run = paragraph.add_run("Table 1. " if lang == "en" else "표 1. ")
    set_run_font(run, lang, 10, True)
    run = paragraph.add_run(caption)
    set_run_font(run, lang, 10)
    paragraph.paragraph_format.keep_with_next = True

    headers = ["Panel", "Method", "Mean F-score", "Geometry loss", "Topology error"]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "D9E2F3")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        set_run_font(run, lang, 8.5, True)

    for row_index, values in enumerate(RESULT_ROWS):
        cells = table.add_row().cells
        for col_index, value in enumerate(values):
            cell = cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if col_index < 2 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            bold = row_index in (0, 3) and col_index >= 2
            run = paragraph.add_run(value)
            set_run_font(run, lang, 8.5, bold)

    content_width = section_content_width_dxa(doc.sections[0])
    widths = column_widths_from_weights([1.18, 2.12, 1.15, 1.15, 1.15], content_width)
    apply_table_geometry(table, widths, table_width_dxa=content_width)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def extract_bibliography(text: str, lang: str) -> list[str]:
    begin = text.find(r"\begin{thebibliography}")
    end = text.find(r"\end{thebibliography}")
    if begin < 0 or end < 0:
        return []
    block = text[begin:end]
    starts = list(re.finditer(r"\\bibitem\{([^{}]+)\}", block))
    result = []
    for index, match in enumerate(starts):
        item_end = starts[index + 1].start() if index + 1 < len(starts) else len(block)
        item = clean_latex(block[match.end() : item_end], lang)
        if lang == "kr":
            item = f"[{index + 1}] {item}"
        result.append(item)
    return result


def add_references(doc: Document, source: str, lang: str, section_number: int) -> None:
    heading_text = (
        f"{section_number} References"
        if lang == "en"
        else f"{section_number} 참고문헌"
    )
    add_text_paragraph(doc, heading_text, lang, "Heading 1")
    for reference in extract_bibliography(source, lang):
        add_text_paragraph(doc, reference, lang, "Bibliography")


def parse_blocks(source: str) -> list[tuple[str, str]]:
    """Return semantic blocks from the manuscript body."""
    body = source[source.find(r"\begin{document}") + len(r"\begin{document}") :]
    body = body[: body.find(r"\end{document}")]
    bibliography_start = body.find(r"\begin{thebibliography}")
    if bibliography_start >= 0:
        body = body[:bibliography_start]

    blocks: list[tuple[str, str]] = []
    pos = 0
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            text = " ".join(line.strip() for line in paragraph_lines).strip()
            paragraph_lines.clear()
            if text:
                blocks.append(("paragraph", text))

    while pos < len(body):
        if body.startswith(r"\begin{abstract}", pos):
            flush_paragraph()
            end = body.find(r"\end{abstract}", pos)
            content = body[pos + len(r"\begin{abstract}") : end]
            blocks.append(("abstract", content))
            pos = end + len(r"\end{abstract}")
            continue
        if body.startswith(r"\begin{figure}", pos):
            flush_paragraph()
            end = body.find(r"\end{figure}", pos)
            blocks.append(("figure", body[pos:end]))
            pos = end + len(r"\end{figure}")
            continue
        if body.startswith(r"\begin{equation}", pos):
            flush_paragraph()
            end = body.find(r"\end{equation}", pos)
            blocks.append(("equation", body[pos:end]))
            pos = end + len(r"\end{equation}")
            continue
        if body.startswith(r"\begin{table}", pos):
            flush_paragraph()
            end = body.find(r"\end{table}", pos)
            blocks.append(("table", body[pos:end]))
            pos = end + len(r"\end{table}")
            continue
        if body.startswith(r"\section{", pos) or body.startswith(r"\subsection{", pos):
            flush_paragraph()
            kind = "section" if body.startswith(r"\section{", pos) else "subsection"
            brace = body.find("{", pos)
            content, end = extract_braced(body, brace)
            blocks.append((kind, content))
            pos = end
            continue

        line_end = body.find("\n", pos)
        if line_end < 0:
            line_end = len(body)
        line = body[pos:line_end].strip()
        pos = line_end + 1
        if not line:
            flush_paragraph()
        elif line in (r"\maketitle", r"\thispagestyle{fancy}"):
            flush_paragraph()
        elif line.startswith("%"):
            continue
        else:
            paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def build_document(
    source_path: Path,
    template_path: Path,
    output_path: Path,
    lang: str,
) -> None:
    source = source_path.read_text(encoding="utf-8")
    doc = Document(str(template_path))
    clear_template_body(doc)
    configure_styles(doc, lang)
    configure_header_footer(doc, lang)

    core = doc.core_properties
    core.title = ""
    core.subject = ""
    core.author = ""
    core.last_modified_by = ""
    core.comments = "Peer-blind manuscript"
    core.keywords = ""
    core.category = ""

    title = clean_latex(extract_command(source, "title"), lang)
    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, lang, 16, True)

    section_number = 0
    subsection_number = 0
    figure_number = 0
    added_keywords = False

    for kind, raw in parse_blocks(source):
        if kind == "abstract":
            add_text_paragraph(
                doc,
                "Abstract" if lang == "en" else "초록",
                lang,
                "Abstract Title",
            )
            add_text_paragraph(doc, clean_latex(raw, lang), lang, "Abstract Body")
            continue
        if kind == "section":
            section_number += 1
            subsection_number = 0
            heading = clean_latex(raw, lang)
            if heading.lower() == "references" or heading == "참고문헌":
                continue
            add_text_paragraph(doc, f"{section_number} {heading}", lang, "Heading 1")
            continue
        if kind == "subsection":
            subsection_number += 1
            heading = clean_latex(raw, lang)
            add_text_paragraph(
                doc,
                f"{section_number}.{subsection_number} {heading}",
                lang,
                "Heading 2",
            )
            continue
        if kind == "equation":
            equation_key = next((key for key in EQUATIONS if f"eq:{key}" in raw), None)
            if equation_key is None:
                equation_key = next((key for key in EQUATIONS if key in raw), None)
            if equation_key:
                add_text_paragraph(doc, EQUATIONS[equation_key], lang, "Equation")
            continue
        if kind == "figure":
            match = re.search(r"\\includegraphics(?:\[[^]]*\])?\{([^{}]+)\}", raw)
            if match:
                image_name = Path(match.group(1)).name
                figure_number += 1
                caption = FIGURE_CAPTIONS[lang][image_name]
                add_figure(
                    doc,
                    source_path.parent / "pics" / image_name,
                    caption,
                    figure_number,
                    lang,
                )
            continue
        if kind == "table":
            add_results_table(doc, lang)
            continue
        if kind == "paragraph":
            if raw.startswith(r"\noindent\textbf{Keywords.") or raw.startswith(
                r"\noindent\textbf{핵심어"
            ):
                text = clean_latex(raw.replace(r"\noindent", ""), lang)
                paragraph = add_text_paragraph(doc, text, lang)
                paragraph.paragraph_format.first_line_indent = Inches(0)
                paragraph.paragraph_format.space_after = Pt(6)
                added_keywords = True
                continue
            text = clean_latex(raw.replace(r"\noindent", ""), lang)
            if text:
                add_text_paragraph(doc, text, lang)

    if not added_keywords:
        raise RuntimeError(f"keywords paragraph not found in {source_path}")
    add_references(doc, source, lang, section_number + 1)

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-root", type=Path, required=True)
    parser.add_argument("--template-docx", type=Path, required=True)
    args = parser.parse_args()

    root = args.project_root.resolve()
    build_document(
        root / "draft" / "paper_en.tex",
        args.template_docx,
        root / "draft" / "paper_en.docx",
        "en",
    )
    build_document(
        root / "draft" / "paper_kr.tex",
        args.template_docx,
        root / "draft" / "paper_kr.docx",
        "kr",
    )
    print(root / "draft" / "paper_en.docx")
    print(root / "draft" / "paper_kr.docx")


if __name__ == "__main__":
    main()
